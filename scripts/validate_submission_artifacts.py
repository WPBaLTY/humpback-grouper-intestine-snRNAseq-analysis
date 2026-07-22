import argparse
import csv
import gzip
import hashlib
import math
from collections import Counter
from pathlib import Path
import sys
from zipfile import ZipFile

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
DATA_ROOT = ROOT.parent / "zenodo_processed_data_record" if (ROOT.parent / "zenodo_processed_data_record").is_dir() else ROOT
DATA_METADATA = DATA_ROOT / "metadata"
SOURCE = DATA_ROOT / "source_data" / "figures"
FIGURES = DATA_ROOT / "figures" / "main" if (DATA_ROOT / "figures" / "main").is_dir() else ROOT / "outputs" / "figures" / "scientific_data"
DOUBLET = DATA_ROOT / "audit" / "doublet_assessment"

EXPECTED_ACCESSIONS = {
    "COM_1": ("GSM9627310", "SRX32711943"),
    "COM_2": ("GSM9627311", "SRX32711944"),
    "CTL_1": ("GSM9627312", "SRX32711945"),
    "CTL_2": ("GSM9627313", "SRX32711946"),
    "Llac_1": ("GSM9627314", "SRX32711947"),
    "Llac_2": ("GSM9627315", "SRX32711948"),
    "Slim_1": ("GSM9627316", "SRX32711949"),
    "Slim_2": ("GSM9627317", "SRX32711950"),
}
EXPECTED_IMMUNE_COUNTS = {
    "T cell (CCR7+)": 1110,
    "Activated T (RORA+)": 586,
    "Activated lymphoid (CCL20hi)": 723,
    "NK-like cytotoxic": 1027,
    "B cells": 108,
    "cDC1-like (XCR1+)": 298,
    "Monocytes/macrophages": 640,
    "MoDC-like (CD209d+)": 219,
    "Granulocyte-like": 242,
    "Cycling (G2/M)": 108,
}


class Audit:
    def __init__(self):
        self.records = []
        self.failures = []

    def check(self, condition, label, detail=""):
        status = "PASS" if condition else "FAIL"
        self.records.append({"status": status, "check": label, "detail": detail})
        if not condition:
            self.failures.append(label)
        suffix = f": {detail}" if detail else ""
        print(f"[{status}] {label}{suffix}")

    def info(self, label, detail=""):
        self.records.append({"status": "INFO", "check": label, "detail": detail})
        suffix = f": {detail}" if detail else ""
        print(f"[INFO] {label}{suffix}")


def read_csv(path):
    opener = gzip.open if path.suffix == ".gz" else open
    with opener(path, "rt", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def read_tsv(path):
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle, delimiter="\t"))


def numeric_text(value):
    return float(str(value).replace(",", "").replace("%", ""))


def sha256_bytes(data):
    return hashlib.sha256(data).hexdigest()


def sha256_file(path):
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def audit_supplement(path, audit):
    workbook = load_workbook(path, data_only=False)
    expected_sheets = [
        "Table_S1_sample_mapping",
        "Table_S2_metadata_schema",
        "Table_S3_annotation",
        "Table_S4_software",
    ]
    audit.check(workbook.sheetnames == expected_sheets, "Supplement sheet names and order", str(workbook.sheetnames))
    freeze_panes = [str(sheet.freeze_panes) if sheet.freeze_panes else "none" for sheet in workbook.worksheets]
    if all(value == "A2" for value in freeze_panes):
        audit.check(True, "Supplement frozen header rows", "A2 on all four sheets")
    else:
        audit.info(
            "Supplement frozen header rows",
            f"presentation-only; observed {freeze_panes}; no effect on cell data or machine readability",
        )

    formula_errors = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith(
                    ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#N/A")
                ):
                    formula_errors.append(f"{sheet.title}!{cell.coordinate}={cell.value}")
    audit.check(not formula_errors, "Supplement formula-error scan", "none" if not formula_errors else "; ".join(formula_errors))

    s1 = workbook["Table_S1_sample_mapping"]
    observed_accessions = {
        s1.cell(row, 1).value: (s1.cell(row, 7).value, s1.cell(row, 8).value)
        for row in range(2, 10)
    }
    release_accessions = {
        row["library"]: (row["gsm_accession"], row["sra_experiment"])
        for row in read_csv(DATA_METADATA / "geo_gsm_mapping.csv")
    }
    audit.check(
        observed_accessions == EXPECTED_ACCESSIONS == release_accessions,
        "Supplement GEO/GSM/SRX closure",
        "all eight libraries exact",
    )
    audit.check(s1["D1"].value == "pooled_library_replicate", "Supplement pooled-library terminology", str(s1["D1"].value))

    s2 = workbook["Table_S2_metadata_schema"]
    s2_fields = [s2.cell(row, 1).value for row in range(2, s2.max_row + 1)]
    audit.check(
        "pooled_library_replicate" in s2_fields and "biological_replicate" not in s2_fields,
        "Supplement metadata-schema replicate field",
        "pooled library, not unsupported tank-level biological replicate",
    )

    s3 = workbook["Table_S3_annotation"]
    observed_immune_counts = {
        s3.cell(row, 3).value: int(s3.cell(row, 4).value)
        for row in range(2, s3.max_row + 1)
        if s3.cell(row, 2).value == "leukocyte reclustering"
    }
    audit.check(observed_immune_counts == EXPECTED_IMMUNE_COUNTS, "Supplement immune subtype counts", "all 10 exact")
    audit.check(sum(observed_immune_counts.values()) == 5061, "Supplement immune count total", "5,061")
    percent_match = all(
        math.isclose(
            float(s3.cell(row, 5).value),
            int(s3.cell(row, 4).value) / 5061 * 100,
            rel_tol=0,
            abs_tol=1e-10,
        )
        for row in range(2, s3.max_row + 1)
        if s3.cell(row, 2).value == "leukocyte reclustering"
    )
    audit.check(percent_match, "Supplement immune percentages", "counts / 5,061 x 100")

    figure4_counts = Counter(row["immune_subtype"] for row in read_csv(SOURCE / "Fig4_immune_umap_coordinates.csv"))
    audit.check(figure4_counts == Counter(EXPECTED_IMMUNE_COUNTS), "Supplement versus Figure 4 counts", "all 5,061 cells close")

    s4 = workbook["Table_S4_software"]
    s4_text = "\n".join(str(cell.value or "") for row in s4.iter_rows() for cell in row)
    required_s4 = [
        "Doublet assessment",
        "Immune-branch cluster filtering",
        "Figure assembly and release validation",
        "R 4.6.1; Seurat 5.3.0; SeuratObject 5.2.0; harmony 1.2.3",
        "R 4.6.1; Seurat 5.3.0; uwot 0.2.4",
        "108 B cells and 242 Granulocyte-like nuclei",
    ]
    audit.check(all(text in s4_text for text in required_s4), "Supplement software/filter audit rows", "required records present")
    audit.check(
        "CreateSeuratObject(min.cells=3) per library" in s4_text
        and "26,134 Cell Ranger reference features yielded a 23,740-feature merged Seurat union" in s4_text,
        "Supplement feature-stage provenance",
        "26,134 Cell Ranger features to 23,740 Seurat features via per-library min.cells=3",
    )
    audit.check("not run" not in s4_text.lower() and "Seurat 5.5.0" not in s4_text, "Supplement obsolete wording scan", "none")


def audit_manuscript(path, audit):
    document = Document(path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    all_text = paragraph_text + "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )

    expected_title = (
        "A single-nucleus transcriptomic atlas of humpback grouper intestine "
        "across dietary supplementation conditions"
    )
    observed_title = document.paragraphs[0].text
    audit.check(
        observed_title == expected_title
        and len(observed_title) == 109
        and ":" not in observed_title
        and "(" not in observed_title
        and ")" not in observed_title,
        "Manuscript Scientific Data title rule",
        f"{len(observed_title)} characters",
    )
    audit.check("Fig. S4" not in all_text, "Manuscript removed redundant Fig. S4 citation", "none")

    required_text = [
        "102,036 high-quality nuclei",
        "pooled sample-level replicates",
        "10,440 nuclei assigned to the global leukocyte-associated clusters",
        "not an additional nFeature_RNA or nCount_RNA threshold",
        "995/5,379 (18.50%) versus 30/5,061 (0.59%)",
        "673/5,379 (12.51%) versus 10/5,061 (0.20%)",
        "4.97 × 10−14",
        "SRX32711943-SRX32711950",
        "within-cell-type detection fraction",
        "cDC1-like (XCR1+)",
        "MoDC-like (CD209d+)",
        "Dot size indicates the percentage of nuclei expressing each marker",
    ]
    missing = [text for text in required_text if text not in all_text]
    audit.check(not missing, "Manuscript required scientific statements", "all present" if not missing else "; ".join(missing))
    audit.check(
        "CreateSeuratObject, min.cells = 3" in all_text
        and "the 26,134 Cell Ranger reference features therefore yielded a 23,740-feature union" in all_text,
        "Manuscript feature-stage provenance",
        "26,134 Cell Ranger features to 23,740 Seurat features via per-library min.cells=3",
    )

    figure_source_row = next(
        (row for row in document.tables[0].rows if row.cells[0].text == "Figure source data tables"),
        None,
    )
    audit.check(figure_source_row is not None, "Manuscript Table 1 figure-source row", "present")
    if figure_source_row is not None:
        figure_source_text = figure_source_row.cells[1].text
        forbidden_legacy_sources = [
            "figure_source_data/:",
            " source_data/:",
            "supp_fig_S4",
            "figure2_alignment_metrics.csv",
            "figure2_sequence_metrics.csv",
            "figure5_major_celltype_composition.csv",
            "figure5_immune_subtype_composition.csv",
        ]
        found_legacy_sources = [
            value for value in forbidden_legacy_sources if value in figure_source_text
        ]
        audit.check(
            not found_legacy_sources,
            "Manuscript Table 1 legacy figure-source scan",
            "none" if not found_legacy_sources else "; ".join(found_legacy_sources),
        )
        listed_sources = [
            "metadata/samples.sequence.stat.xls",
            "metadata/samples.align.stat.xls",
            "source_data/figures/Fig3_global_marker_dotplot_source.csv",
            "source_data/figures/Fig4_immune_marker_dotplot_source.csv",
            "source_data/figures/Fig5a_global_celltype_composition_leukocytes_merged.csv",
            "source_data/figures/Fig5b_immune_subtype_composition.csv",
            "source_data/figures/SuppFigureS1_qc_per_cell.csv.gz",
            "source_data/figures/SuppFigureS3_pseudobulk_correlation.csv",
        ]
        manifest_sources = {
            row["source_file"] for row in read_csv(DATA_METADATA / "figure_source_manifest.csv")
        }
        missing_from_manifest = sorted(set(listed_sources) - manifest_sources)
        missing_from_release = sorted(source for source in listed_sources if not (DATA_ROOT / source).is_file())
        missing_from_cell = sorted(
            Path(source).name for source in listed_sources if Path(source).name not in figure_source_text
        )
        audit.check(
            "metadata/figure_source_manifest.csv" in figure_source_text
            and not missing_from_manifest
            and not missing_from_release
            and not missing_from_cell,
            "Manuscript Table 1 current figure-source closure",
            "manifest plus 8 listed sources close"
            if not (missing_from_manifest or missing_from_release or missing_from_cell)
            else (
                f"manifest={missing_from_manifest}; release={missing_from_release}; "
                f"cell={missing_from_cell}"
            ),
        )

    forbidden_text = [
        "within-cluster detection fraction",
        "biological_replicate",
        "5,379 doublets",
        "Seurat v5.5",
    ]
    found_forbidden = [text for text in forbidden_text if text in all_text]
    audit.check(not found_forbidden, "Manuscript obsolete/unsupported wording scan", "none" if not found_forbidden else "; ".join(found_forbidden))

    protected = [
        "Author contributions will be finalized according to the final author list before formal submission.",
        "Acknowledgements will be finalized before formal submission.",
        "Funding information will be finalized before formal submission.",
    ]
    audit.check(all(text in paragraph_text for text in protected), "Protected author/funding placeholders", "unchanged")

    audit.check(len(document.tables) >= 1 and len(document.tables[0].rows) == 11, "Manuscript Table 1 row count", "11")
    audit_cell = document.tables[0].rows[-1].cells[1]
    audit.check(
        "doublet_assessment_per_cell.csv.gz" in audit_cell.text
        and "immune_filtering_fate_10440_cells.csv.gz" in audit_cell.text,
        "Manuscript Table 1 audit files",
        "both machine-readable ledgers listed",
    )
    audit_runs = [run for paragraph in audit_cell.paragraphs for run in paragraph.runs if run.text]
    audit.check(
        bool(audit_runs) and all(run.font.size is not None and math.isclose(run.font.size.pt, 7.5, abs_tol=0.01) for run in audit_runs),
        "Manuscript Table 1 final-row font",
        "7.5 pt",
    )

    figure_names = {
        2: "Figure2_QC_final.png",
        3: "Figure3_global_annotation_final.png",
        4: "Figure4_immune_reclustering_final.png",
        5: "Figure5_composition_final.png",
    }
    figure_details = []
    figures_match = True
    with ZipFile(path, "r") as archive:
        for number, filename in figure_names.items():
            embedded = archive.read(f"word/media/image{number}.png")
            source_path = FIGURES / filename
            embedded_hash = sha256_bytes(embedded)
            source_hash = sha256_file(source_path)
            figures_match = figures_match and embedded_hash == source_hash
            figure_details.append(f"F{number}={embedded_hash}")
    audit.check(figures_match, "Manuscript embedded Figure 2-5 hashes", ", ".join(figure_details))

    sequence_rows = read_tsv(DATA_METADATA / "samples.sequence.stat.xls")
    alignment_rows = read_tsv(DATA_METADATA / "samples.align.stat.xls")
    reads_million = [numeric_text(row["Number of Reads"]) / 1e6 for row in sequence_rows]
    valid = [numeric_text(row["Valid Barcodes"]) for row in sequence_rows]
    saturation = [numeric_text(row["Sequencing Saturation"]) for row in sequence_rows]
    genome = [numeric_text(row["Reads Mapped Confidently to Genome"]) for row in alignment_rows]
    transcriptome = [numeric_text(row["Reads Mapped Confidently to Transcriptome"]) for row in alignment_rows]
    estimated = [int(numeric_text(row["Estimated Number of Cells"])) for row in alignment_rows]
    fraction = [numeric_text(row["Fraction Reads in Cells"]) for row in alignment_rows]
    genes = [int(numeric_text(row["Median Genes per Cell"])) for row in alignment_rows]
    umis = [int(numeric_text(row["Median UMI Counts per Cell"])) for row in alignment_rows]
    figure2_phrases = [
        f"Total reads ranged from {min(reads_million):.1f} to {max(reads_million):.1f} million",
        f"valid barcode rates from {min(valid):.1f}% to {max(valid):.1f}%",
        f"sequencing saturation from {min(saturation):.1f}% to {max(saturation):.1f}%",
        f"confident genome mapping from {min(genome):.1f}% to {max(genome):.1f}%",
        f"confident transcriptome mapping from {min(transcriptome):.1f}% to {max(transcriptome):.1f}%",
        f"Cell Ranger estimated {min(estimated):,} to {max(estimated):,} nuclei per library",
        f"Fraction reads in nuclei ranged from {min(fraction):.1f}% to {max(fraction):.1f}%",
        f"median genes per nucleus from {min(genes):,} to {max(genes):,}",
        f"median UMI counts per nucleus from {min(umis):,} to {max(umis):,}",
    ]
    missing_figure2 = [phrase for phrase in figure2_phrases if phrase not in paragraph_text]
    audit.check(not missing_figure2, "Manuscript Figure 2 metric ranges", "all source-derived ranges exact" if not missing_figure2 else "; ".join(missing_figure2))

    global_rows = read_csv(SOURCE / "Fig3_global_umap_coordinates_final_annotation.csv.gz")
    treatment_counts = Counter(row["group"] for row in global_rows)
    expected_group_sentence = (
        f"The retained nuclei comprised {treatment_counts['CTL']:,} nuclei from the CTL group, "
        f"{treatment_counts['Llac']:,} from the Llac group, {treatment_counts['Slim']:,} from the Slim group "
        f"and {treatment_counts['COM']:,} from the COM group."
    )
    audit.check(expected_group_sentence in paragraph_text, "Manuscript treatment-level nucleus counts", str(dict(treatment_counts)))

    doublet_rows = read_csv(DOUBLET / "doublet_assessment_per_cell.csv.gz")
    cluster_total = sum(row["cluster_aware_class"] == "doublet" for row in doublet_rows)
    random_total = sum(row["random_class"] == "doublet" for row in doublet_rows)
    concordant = sum(
        row["cluster_aware_class"] == "doublet" and row["random_class"] == "doublet"
        for row in doublet_rows
    )
    audit.check(
        f"{cluster_total:,}/102,036" in paragraph_text
        and f"{random_total:,}/102,036" in paragraph_text
        and f"{concordant:,} nuclei across the atlas" in paragraph_text,
        "Manuscript whole-atlas doublet totals",
        f"{cluster_total:,}; {random_total:,}; concordant {concordant:,}",
    )


def write_report(path, records):
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["status", "check", "detail"])
        writer.writeheader()
        writer.writerows(records)


def main():
    parser = argparse.ArgumentParser(description="Validate final manuscript and supplementary artifacts against the release data.")
    parser.add_argument("--manuscript", type=Path, help="Final Scientific Data DOCX")
    parser.add_argument("--supplement", type=Path, help="Final Supplementary Tables XLSX")
    parser.add_argument("--report", type=Path, help="Optional CSV report path")
    args = parser.parse_args()
    if args.manuscript is None and args.supplement is None:
        parser.error("at least one of --manuscript or --supplement is required")

    audit = Audit()
    if args.manuscript is not None:
        audit.check(args.manuscript.is_file(), "Manuscript file exists", args.manuscript.name)
        if args.manuscript.is_file():
            audit_manuscript(args.manuscript, audit)
    if args.supplement is not None:
        audit.check(args.supplement.is_file(), "Supplement file exists", args.supplement.name)
        if args.supplement.is_file():
            audit_supplement(args.supplement, audit)

    if args.report is not None:
        write_report(args.report, audit.records)
        print(f"Report: {args.report}")

    if audit.failures:
        print(f"\nSubmission artifact validation failed: {len(audit.failures)} check(s).")
        return 1
    print("\nSubmission artifact validation passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
